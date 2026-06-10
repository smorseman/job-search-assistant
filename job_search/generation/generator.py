"""Document generation — tailored resume + cover letter via Claude API."""

from __future__ import annotations

import json
import logging
from pathlib import Path

import anthropic
import yaml
from docx import Document as DocxDocument

from job_search.config import settings
from job_search.models import CanonicalJob

from .keywords import KeywordExtractor

logger = logging.getLogger(__name__)

# Prefer claude-sonnet-4-6 with prompt caching for cost efficiency on repeated profile loads
MODEL = "claude-sonnet-4-6"


RESUME_SYSTEM = """You are an expert civil engineering resume writer.
Your task is to produce a tailored, ATS-optimized resume from the candidate's master profile.

RULES (non-negotiable):
1. Every claim must trace to a verified fact in the master profile. No fabrication.
2. Do not inflate responsibilities, overstate proficiency, or imply unearned qualifications.
3. Return ONLY valid JSON matching the schema below. No markdown, no prose outside the schema.
4. Do not include a cover letter — that is a separate call.

OUTPUT SCHEMA:
{
  "professional_summary": "<2-4 sentence positioning statement for THIS role>",
  "skills": ["<skill>", ...],
  "education": [{"degree": "", "major": "", "institution": "", "graduation": "", "gpa": null, "honors": [], "relevant_coursework": []}],
  "experience": [{"employer": "", "title": "", "dates": "", "location": "", "bullets": ["..."]}],
  "projects": [{"name": "", "role": "", "date": "", "bullets": ["..."]}],
  "activities": [{"organization": "", "role": "", "dates": "", "bullets": ["..."]}],
  "certifications": ["<cert or status>"],
  "keyword_notes": {"included": ["..."], "rationale": "<brief>"}
}"""

COVER_LETTER_SYSTEM = """You are an expert civil engineering career advisor.
Write a tailored cover letter for the candidate applying to this specific role.

RULES:
1. Every claim traces to the master profile. No fabrication.
2. The letter is human-facing — professional, readable, not keyword-stuffed.
3. Weave in ONLY top-tier JD keywords naturally; do not pad for density.
4. The letter is NOT a restatement of the resume. It positions WHY this firm, WHY this role.
5. Return ONLY valid JSON: {"salutation": "...", "body_paragraphs": ["...", "..."], "closing": "..."}
6. 3-4 body paragraphs maximum. Each paragraph = one string in the array."""


class DocumentGenerator:
    def __init__(self):
        self.client = anthropic.Anthropic(api_key=settings.ANTHROPIC_API_KEY)
        self.extractor = KeywordExtractor()
        self._profile: dict | None = None

    def load_profile(self, path: str | None = None) -> None:
        profile_path = path or settings.PROFILE_PATH
        with open(profile_path) as f:
            self._profile = yaml.safe_load(f)
        logger.info("Profile loaded from %s", profile_path)

    def generate(self, job: CanonicalJob) -> dict:
        """Generate tailored resume + cover letter for a job. Returns metadata dict."""
        if not self._profile:
            self.load_profile()

        jd = job.description_normalized or job.description_raw or ""
        keywords = self.extractor.extract(jd)

        resume_json = self._generate_resume(job, jd, keywords)
        cover_json = self._generate_cover_letter(job, jd, keywords, resume_json)

        resume_text = self._render_resume_text(resume_json)
        cover_text = self._render_cover_text(cover_json)

        coverage, hits, misses = self.extractor.compute_coverage(keywords, resume_text)

        return {
            "resume_json": resume_json,
            "cover_letter_json": cover_json,
            "resume_text": resume_text,
            "cover_letter_text": cover_text,
            "keyword_coverage": coverage,
            "keywords_hit": hits,
            "keywords_missed": misses,
        }

    def _generate_resume(self, job: CanonicalJob, jd: str, keywords) -> dict:
        profile_json = json.dumps(self._profile, indent=2)
        kw_summary = ", ".join(k.keyword for k in keywords if k.tier == 1)

        messages = [
            {
                "role": "user",
                "content": f"""MASTER PROFILE:
{profile_json}

TARGET ROLE:
Company: {job.company}
Title: {job.title}
Location: {job.location_city}, {job.location_state}
ATS Type: {job.ats_type.value}

JOB DESCRIPTION:
{jd[:8000]}

TIER-1 KEYWORDS (must appear in resume, both long-form and acronym):
{kw_summary}

Produce the tailored resume JSON. Emphasize the most relevant experience and projects for this discipline.
The professional summary must be rewritten specifically for this role.""",
            }
        ]

        resp = self.client.messages.create(
            model=MODEL,
            max_tokens=4096,
            system=[
                {
                    "type": "text",
                    "text": RESUME_SYSTEM,
                    "cache_control": {"type": "ephemeral"},  # cache the system prompt
                }
            ],
            messages=messages,
        )

        raw = resp.content[0].text.strip()
        return json.loads(raw)

    def _generate_cover_letter(
        self,
        job: CanonicalJob,
        jd: str,
        keywords,
        resume_json: dict,
    ) -> dict:
        profile_json = json.dumps(self._profile, indent=2)
        top_keywords = ", ".join(k.keyword for k in keywords if k.tier == 1)

        messages = [
            {
                "role": "user",
                "content": f"""MASTER PROFILE:
{profile_json}

TAILORED RESUME ALREADY GENERATED (use for consistency — do not repeat):
{json.dumps(resume_json, indent=2)[:3000]}

TARGET ROLE:
Company: {job.company}
Title: {job.title}
Location: {job.location_city}, {job.location_state}

JOB DESCRIPTION:
{jd[:6000]}

TOP KEYWORDS (weave in naturally, not mechanically):
{top_keywords}

Write the cover letter JSON. Make it compelling for a human engineering hiring manager.
3-4 paragraphs. Do not summarize the resume — make the case for fit.""",
            }
        ]

        resp = self.client.messages.create(
            model=MODEL,
            max_tokens=2048,
            system=[
                {
                    "type": "text",
                    "text": COVER_LETTER_SYSTEM,
                    "cache_control": {"type": "ephemeral"},
                }
            ],
            messages=messages,
        )

        raw = resp.content[0].text.strip()
        return json.loads(raw)

    def _render_resume_text(self, data: dict) -> str:
        """Plain-text render of the resume JSON for keyword coverage computation."""
        parts = [data.get("professional_summary", "")]
        parts += data.get("skills", [])
        for exp in data.get("experience", []):
            parts.append(exp.get("title", ""))
            parts += exp.get("bullets", [])
        for proj in data.get("projects", []):
            parts.append(proj.get("name", ""))
            parts += proj.get("bullets", [])
        for edu in data.get("education", []):
            parts += edu.get("relevant_coursework", [])
        parts += data.get("certifications", [])
        return " ".join(parts)

    def _render_cover_text(self, data: dict) -> str:
        parts = [data.get("salutation", "")]
        parts += data.get("body_paragraphs", [])
        parts.append(data.get("closing", ""))
        return " ".join(parts)

    def save_docx(self, resume_json: dict, output_path: str) -> None:
        """Save the resume as a clean single-column .docx for ATS parsing."""
        doc = DocxDocument()
        profile = self._profile or {}
        identity = profile.get("identity", {})

        name = f"{identity.get('first_name', '')} {identity.get('last_name', '')}".strip()
        if name:
            heading = doc.add_heading(name, level=0)

        contact_parts = [x for x in [
            identity.get("email"),
            identity.get("phone"),
            identity.get("linkedin_url"),
            f"{identity.get('location', {}).get('city')}, {identity.get('location', {}).get('state')}",
        ] if x]
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

        summary = resume_json.get("professional_summary")
        if summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(summary)

        skills = resume_json.get("skills", [])
        if skills:
            doc.add_heading("Technical Skills", level=1)
            doc.add_paragraph(", ".join(skills))

        for exp in resume_json.get("experience", []):
            doc.add_heading("Experience", level=1) if exp == resume_json["experience"][0] else None
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('title', '')} — {exp.get('employer', '')}").bold = True
            doc.add_paragraph(f"{exp.get('dates', '')} | {exp.get('location', '')}")
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

        for proj in resume_json.get("projects", []):
            doc.add_heading("Projects", level=1) if proj == resume_json["projects"][0] else None
            p = doc.add_paragraph()
            p.add_run(f"{proj.get('name', '')} — {proj.get('role', '')} ({proj.get('date', '')})").bold = True
            for bullet in proj.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

        for edu in resume_json.get("education", []):
            doc.add_heading("Education", level=1) if edu == resume_json["education"][0] else None
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')} in {edu.get('major', '')}").bold = True
            doc.add_paragraph(f"{edu.get('institution', '')} — {edu.get('graduation', '')}")
            if edu.get("relevant_coursework"):
                doc.add_paragraph("Relevant Coursework: " + ", ".join(edu["relevant_coursework"]))

        certs = resume_json.get("certifications", [])
        if certs:
            doc.add_heading("Certifications & Licensure", level=1)
            for cert in certs:
                doc.add_paragraph(cert, style="List Bullet")

        doc.save(output_path)
        logger.info("Resume saved to %s", output_path)
